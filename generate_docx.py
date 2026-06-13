import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set cell background color (hex format, e.g. 'F8FAFC')"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='CBD5E1', sz='4', val='single'):
    """Set thin border around the cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
        
    tcPr.append(tcBorders)

def add_header_banner(doc, title_text, subtitle_text):
    # Setup page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("SIENVI AGENCY")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(148, 163, 184) # Slate 400
    title_p.paragraph_format.space_after = Pt(2)
    
    main_title_p = doc.add_paragraph()
    main_title_run = main_title_p.add_run(title_text)
    main_title_run.font.name = 'Arial'
    main_title_run.font.size = Pt(20)
    main_title_run.font.bold = True
    main_title_run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    main_title_p.paragraph_format.space_after = Pt(6)
    
    sub_title_p = doc.add_paragraph()
    sub_title_run = sub_title_p.add_run(subtitle_text)
    sub_title_run.font.name = 'Arial'
    sub_title_run.font.size = Pt(10)
    sub_title_run.font.italic = True
    sub_title_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    sub_title_p.paragraph_format.space_after = Pt(18)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue 600

def add_question(doc, label, description=None, is_textarea=False, options=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    
    run_label = p.add_run(label)
    run_label.font.name = 'Arial'
    run_label.font.size = Pt(10)
    run_label.font.bold = True
    run_label.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    
    if description:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(3)
        p_desc.paragraph_format.keep_with_next = True
        run_desc = p_desc.add_run(description)
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(9)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(148, 163, 184) # Slate 400
        
    if options:
        for opt in options:
            p_opt = doc.add_paragraph(style='List Bullet')
            p_opt.paragraph_format.space_before = Pt(0)
            p_opt.paragraph_format.space_after = Pt(2)
            
            run_opt_box = p_opt.add_run("☐  ")
            run_opt_box.font.name = 'Arial'
            run_opt_box.font.size = Pt(10)
            
            run_opt_text = p_opt.add_run(opt)
            run_opt_text.font.name = 'Arial'
            run_opt_text.font.size = Pt(10)
            run_opt_text.font.color.rgb = RGBColor(71, 85, 105)
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(4)
    else:
        # Create an input box table (1x1 table)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        
        # Set background fill (F8FAFC - Slate 50) and borders (CBD5E1 - Slate 300)
        set_cell_background(cell, 'F8FAFC')
        set_cell_borders(cell, 'CBD5E1', sz='4')
        
        # Set box height / spacing
        cell_p = cell.paragraphs[0]
        cell_p.paragraph_format.space_before = Pt(4)
        if is_textarea:
            # Leave space for multiple lines of typing
            cell_p.paragraph_format.space_after = Pt(40)
        else:
            cell_p.paragraph_format.space_after = Pt(12)
            
        run_placeholder = cell_p.add_run("Type your answer here...")
        run_placeholder.font.name = 'Arial'
        run_placeholder.font.size = Pt(9.5)
        run_placeholder.font.italic = True
        run_placeholder.font.color.rgb = RGBColor(203, 213, 225) # Slate 300
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_business_admin_docx(output_path):
    doc = docx.Document()
    add_header_banner(doc, "Non-Ecommerce / Business Admin Client Onboarding Questionnaire", 
                      "Purpose: This questionnaire helps us understand your business, workflows, administrative needs, bottlenecks, and success goals so we can support you efficiently. Estimated completion: 30–60 minutes.")
    
    # Section 1
    add_section_header(doc, "Section 1: Business & Contact Information")
    add_question(doc, "Business Name *", "Legal and/or operating business name.")
    add_question(doc, "Primary Contact Name *", "Name of the person coordinating with Sienvi Agency.")
    add_question(doc, "Role / Title *", "Founder, CEO, Operations Manager, Executive Assistant, etc.")
    add_question(doc, "Email Address *")
    add_question(doc, "Phone Number / WhatsApp Number")
    add_question(doc, "Business Website")
    add_question(doc, "Social Media / LinkedIn Links")
    add_question(doc, "Business Location / Time Zone *")
    add_question(doc, "Preferred Communication Method *", "Email, Slack, WhatsApp, ClickUp, Asana, Monday, Zoom, Google Chat, other.")
    add_question(doc, "Preferred Meeting Frequency", "Weekly, bi-weekly, monthly, as-needed, no recurring meetings.")

    # Section 2
    add_section_header(doc, "Section 2: Business Overview")
    add_question(doc, "How would you describe your business? *", "What do you do, who do you serve, and what is your main value proposition?", is_textarea=True)
    add_question(doc, "Industry / Niche *", "Examples: consulting, coaching, real estate, healthcare, legal, creative agency, nonprofit, professional services.")
    add_question(doc, "How long have you been operating?")
    add_question(doc, "What are your primary services or offers? *", "List your main services, packages, retainers, programs, or client deliverables.", is_textarea=True)
    add_question(doc, "Who are your primary clients or customers? *", "Briefly describe the types of people or businesses you serve.")
    add_question(doc, "What are your current revenue streams?", "Examples: retainers, hourly services, project work, subscriptions, memberships, commissions, licensing.")
    add_question(doc, "What stage is your business currently in?", "Startup, growing, scaling, stabilizing, restructuring, overwhelmed chaos goblin mode — professionally speaking.")

    # Section 3
    add_section_header(doc, "Section 3: Goals & Success Outcomes")
    add_question(doc, "What are your top 3 business goals for the next 6–12 months? *", "List goals 1, 2, and 3 here.", is_textarea=True)
    add_question(doc, "What would a successful partnership with Sienvi Agency look like? *", "Be specific. Examples: fewer missed follow-ups, cleaner inbox, organized CRM, smoother client onboarding, documented SOPs, faster invoicing, better reporting.", is_textarea=True)
    add_question(doc, "What specific outcomes do you want from business-admin support? *", "Examples: save 10 hours/week, reduce missed emails, organize files, manage calendar, streamline client communication.", is_textarea=True)
    add_question(doc, "How will we measure success? *", "Examples: response time, task completion rate, fewer overdue items, number of SOPs created, inbox zero, lead follow-up speed, reporting accuracy.", is_textarea=True)
    add_question(doc, "Do you have any deadlines, launches, busy seasons, or urgent priorities?")
    add_question(doc, "Long-term vision: Where do you want the business to be in 3–5 years?")

    # Section 4
    add_section_header(doc, "Section 4: Current Challenges & Bottlenecks")
    add_question(doc, "What are the biggest administrative challenges in your business right now? *", "Describe what feels messy, slow, repetitive, or stressful.", is_textarea=True)
    add_question(doc, "Where are you personally spending too much time? *", "Examples: emails, scheduling, client follow-ups, invoicing, file organization, task tracking, reporting, data entry.", is_textarea=True)
    add_question(doc, "What tasks are falling through the cracks?")
    add_question(doc, "What processes feel unclear, undocumented, or inconsistent?")
    add_question(doc, "What is currently preventing your business from operating smoothly?")
    add_question(doc, "Have you worked with admin support, VAs, OBMs, or agencies before? What worked? What did not?", is_textarea=True)
    add_question(doc, "What would you like to stop doing yourself immediately? *", "No heroic founder syndrome here. If a $15/hour task is stealing CEO energy, we are sending it to the delegation graveyard.", is_textarea=True)

    # Section 5
    add_section_header(doc, "Section 5: Admin Support Needs")
    add_question(doc, "Which business-admin services do you need help with? Select all that apply. *", None, options=[
        "Inbox management", "Calendar management", "Appointment scheduling", "Client onboarding", "Client offboarding",
        "CRM management", "Lead tracking", "Follow-up management", "Data entry", "Document formatting", "File organization",
        "SOP creation", "Project management", "Meeting notes", "Reporting / dashboards", "Invoice support", "Payment follow-ups",
        "Contract / proposal support", "Research", "Vendor coordination", "Team coordination", "Travel planning",
        "Personal executive support"
    ])
    add_question(doc, "Which tasks are highest priority right now? *", "List the top 3–5.", is_textarea=True)
    add_question(doc, "Which tasks are recurring?", "Daily, weekly, monthly, quarterly.")
    add_question(doc, "Which tasks are one-time cleanup projects?")
    add_question(doc, "Are there any tasks we should NOT handle?", "Legal decisions, financial approvals, sensitive HR issues, final client decisions, etc.")

    # Section 6
    add_section_header(doc, "Section 6: Workflow & Operations")
    add_question(doc, "Describe your current workflow from lead/client inquiry to completed service. *", "Walk us through what happens step by step.", is_textarea=True)
    add_question(doc, "Do you have documented SOPs or processes? *", None, options=["Yes", "No", "Partially"])
    add_question(doc, "Where are your SOPs stored?", "Google Drive, Notion, ClickUp, Loom, Trainual, Dropbox, other.")
    add_question(doc, "What processes need to be documented first?")
    add_question(doc, "How are tasks currently assigned and tracked? *", "ClickUp, Asana, Monday, Trello, Notion, spreadsheets, email, memory and prayer.")
    add_question(doc, "What recurring meetings or check-ins does your business currently have?")
    add_question(doc, "Are there approval steps we need to follow before completing or sending anything?")
    add_question(doc, "Who makes final decisions on admin, client, finance, or operations items?")

    # Section 7
    add_section_header(doc, "Section 7: Tools, Platforms & Access")
    add_question(doc, "Communication Tools Used", None, options=["Slack", "Email", "WhatsApp", "Microsoft Teams", "Google Chat"])
    add_question(doc, "Project Management Tools Used", None, options=["ClickUp", "Asana", "Trello", "Monday.com", "Notion", "Airtable"])
    add_question(doc, "CRM Tools Used", None, options=["HubSpot", "GoHighLevel", "Salesforce", "Pipedrive", "Dubsado", "HoneyBook", "Spreadsheets"])
    add_question(doc, "File Storage Tools Used", None, options=["Google Drive", "Dropbox", "OneDrive", "Notion"])
    add_question(doc, "Finance / Invoicing Tools Used", None, options=["QuickBooks Online", "Xero", "Wave", "Stripe", "PayPal", "FreshBooks", "HoneyBook"])
    add_question(doc, "What tools need to be cleaned up, organized, or connected?", is_textarea=True)
    add_question(doc, "How will access be granted?", "Password manager, direct invite, shared admin account, temporary access.")
    add_question(doc, "Password Manager Used", "1Password, LastPass, Dashlane, Google Password Manager, none.")

    # Section 8
    add_section_header(doc, "Section 8: Client / Customer Communication")
    add_question(doc, "Who do you communicate with regularly?", "Clients, leads, vendors, team members, partners, contractors.")
    add_question(doc, "What types of messages do you receive most often?", "New inquiries, support questions, scheduling requests, billing questions, project updates, complaints, referrals.")
    add_question(doc, "Do you have email templates or message scripts?", None, options=["Yes", "No", "Partially"])
    add_question(doc, "What tone should we use when communicating on your behalf? *", "Professional, warm, friendly, concise, premium, corporate, casual, authoritative, other.")
    add_question(doc, "Are there phrases, promises, or claims we should avoid?")
    add_question(doc, "What situations require your approval before responding? *", "Examples: refunds, discounts, legal topics, angry clients, high-value leads, partnership opportunities.", is_textarea=True)
    add_question(doc, "What is your desired response time for client messages?", "Same day, within 24 hours, within 48 hours, urgent only.")

    # Section 9
    add_section_header(doc, "Section 9: Sales, Leads & Follow-Up")
    add_question(doc, "How do new leads currently come into the business?", "Website form, referrals, social media, email, phone, ads, networking, events, partnerships.")
    add_question(doc, "Where are leads tracked?")
    add_question(doc, "What happens after someone becomes a lead?", "Describe the current follow-up process.", is_textarea=True)
    add_question(doc, "Do you have a sales pipeline or CRM stages?", "Example: New Lead -> Contacted -> Discovery Call Booked -> Proposal Sent -> Closed Won/Lost.")
    add_question(doc, "Where do leads commonly fall through the cracks?")
    add_question(doc, "Do you need help with lead follow-ups?")
    add_question(doc, "Do you need help preparing proposals, contracts, or onboarding materials?")
    add_question(doc, "What qualifies a good lead for your business?")
    add_question(doc, "Are there clients or leads you want to avoid?", "Red flags, poor fit, low budget, high maintenance, wrong niche.")

    # Section 10
    add_section_header(doc, "Section 10: Calendar & Scheduling")
    add_question(doc, "Who manages your calendar currently?")
    add_question(doc, "What types of appointments need scheduling?", "Sales calls, client calls, internal meetings, interviews, vendor meetings, personal appointments.")
    add_question(doc, "What are your working hours and availability rules?")
    add_question(doc, "Are there protected focus times or no-meeting days?")
    add_question(doc, "What meetings require prep materials?")
    add_question(doc, "What meetings require follow-up notes or action items?")
    add_question(doc, "Do you use scheduling links? If yes, provide them.")
    add_question(doc, "Any scheduling preferences we should know?", "Buffer time, maximum calls per day, preferred days, time zone rules.")

    # Section 11
    add_section_header(doc, "Section 11: Documents, Files & Organization")
    add_question(doc, "Where are your business files currently stored? *")
    add_question(doc, "How organized are your files right now?", "Very organized, somewhat organized, messy, digital junk drawer.")
    add_question(doc, "What file categories do you use or need?", "Clients, finance, contracts, marketing, operations, team, templates, SOPs, personal, other.")
    add_question(doc, "Do you need a new folder structure created?")
    add_question(doc, "Which documents do you use repeatedly?", "Proposals, contracts, invoices, onboarding forms, templates.")
    add_question(doc, "Do any documents need formatting, updating, or standardization?")
    add_question(doc, "Are there naming conventions we should follow?")

    # Section 12
    add_section_header(doc, "Section 12: Reporting & Performance Tracking")
    add_question(doc, "What numbers or KPIs do you currently track?", "Revenue, leads, calls booked, client retention, project deadlines, task completion, cash collected, email response time.")
    add_question(doc, "Where do you track performance?", "Spreadsheet, CRM, dashboard, accounting software, project management tool.")
    add_question(doc, "What reports do you need created or maintained?", "Weekly summary, monthly revenue report, lead report, client status report, project progress report, admin task report.")
    add_question(doc, "Who should receive reports?")
    add_question(doc, "How often should reports be sent?", "Daily, weekly, bi-weekly, monthly.")
    add_question(doc, "What does leadership need to know at a glance?")

    # Section 13
    add_section_header(doc, "Section 13: Finance & Billing Support")
    add_question(doc, "Do you need support with invoicing or payment follow-up?")
    add_question(doc, "Which finance tools do you use?")
    add_question(doc, "Who approves invoices, expenses, refunds, or payment reminders?")
    add_question(doc, "Do you have invoice templates or payment reminder scripts?")
    add_question(doc, "How often are invoices sent?")
    add_question(doc, "How should overdue payments be handled?")
    add_question(doc, "Are there financial tasks we should not touch?")

    # Section 14
    add_section_header(doc, "Section 14: Team & Vendor Coordination")
    add_question(doc, "Do you have a team?", "List team members, roles, and responsibilities.", is_textarea=True)
    add_question(doc, "Do you work with contractors or vendors?")
    add_question(doc, "Who needs to be included in communication or project updates?")
    add_question(doc, "Are there recurring team tasks we should help coordinate?")
    add_question(doc, "Do you need help assigning tasks or following up with team members?")
    add_question(doc, "Are there internal communication rules we should follow?")

    # Section 15
    add_section_header(doc, "Section 15: Brand, Voice & Professional Standards")
    add_question(doc, "Do you have brand guidelines?", "Logo, colors, fonts, tone of voice, templates.")
    add_question(doc, "How should your business sound in written communication? *")
    add_question(doc, "Are there examples of communication you like?", "Emails, proposals, client messages, newsletters, reports.")
    add_question(doc, "Are there examples you dislike or want to avoid?")
    add_question(doc, "What standards matter most to you?", "Speed, accuracy, warmth, detail, professionalism, confidentiality, initiative, minimal back-and-forth.")
    add_question(doc, "Any legal, compliance, privacy, or confidentiality requirements?")

    # Section 16
    add_section_header(doc, "Section 16: Access, Security & Confidentiality")
    add_question(doc, "What systems will we need access to? *", is_textarea=True)
    add_question(doc, "Who will send access invitations?")
    add_question(doc, "Do you require NDA or confidentiality agreements?")
    add_question(doc, "Are there sensitive files, clients, or accounts with restricted access?")
    add_question(doc, "Do you have two-factor authentication requirements?")
    add_question(doc, "What is the process for revoking access if needed?")
    add_question(doc, "Any data privacy rules we must follow?")

    # Section 17
    add_section_header(doc, "Section 17: Collaboration & Feedback")
    add_question(doc, "Who is the final decision-maker? *")
    add_question(doc, "Who approves completed work?")
    add_question(doc, "How do you prefer to give feedback?", "Loom, comments, email, Slack, live call, project management tool.")
    add_question(doc, "What turnaround time do you expect for feedback from your side?")
    add_question(doc, "What turnaround time do you expect from us?")
    add_question(doc, "How should urgent requests be handled?")
    add_question(doc, "What does your ideal working relationship look like?")
    add_question(doc, "What frustrates you most when working with service providers?")

    # Section 18
    add_section_header(doc, "Section 18: Priorities & First 30 Days")
    add_question(doc, "What should we fix first? *", "Pick the biggest operational pain.", is_textarea=True)
    add_question(doc, "What should we complete in the first 7 days?")
    add_question(doc, "What should we complete in the first 30 days?")
    add_question(doc, "What would make you say, \"This was absolutely worth it\" after the first month? *", is_textarea=True)
    add_question(doc, "Are there any immediate fires we need to put out?")
    add_question(doc, "Are there any upcoming deadlines we need to know about?")

    # Section 19
    add_section_header(doc, "Section 19: Assets & Links")
    add_question(doc, "Please enter Google Drive / PM / CRM URLs here:", "Drive Folder, PM Workspace, CRM invite, Scheduling, Website, Brand assets...", is_textarea=True)

    # Section 20
    add_section_header(doc, "Section 20: Final Notes")
    add_question(doc, "Final comments, outsourcing worries, or preferences?", is_textarea=True)

    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

def build_general_discovery_docx(output_path):
    doc = docx.Document()
    add_header_banner(doc, "General Client Onboarding & Discovery Questionnaire", 
                      "Purpose: This form helps us understand your business, goals, audience, offers, challenges, current systems, and project needs so we can create a clear strategy and execute with precision. Estimated completion: 30–60 minutes.")

    # Section 1
    add_section_header(doc, "Section 1: Basic Business Information")
    add_question(doc, "Business Name *", "Your legal or operating business name.")
    add_question(doc, "Primary Contact Name *")
    add_question(doc, "Role / Title *")
    add_question(doc, "Email Address *")
    add_question(doc, "Phone Number / WhatsApp Number")
    add_question(doc, "Website URL")
    add_question(doc, "Social Media Links", "Instagram, Facebook, LinkedIn, TikTok, YouTube, Pinterest, etc.")
    add_question(doc, "Business Location / Time Zone *")
    add_question(doc, "Preferred Communication Method *", "Email, Slack, WhatsApp, ClickUp, Asana, Monday, Zoom, Google Chat, other.")
    add_question(doc, "Who will be the main point of contact? *", "Name, role, and email.")
    add_question(doc, "Are you the final decision-maker? *", None, options=["Yes", "No"])

    # Section 2
    add_section_header(doc, "Section 2: Business Overview")
    add_question(doc, "How would you describe your business? *", "Explain what you do, who you serve, and how you help them.", is_textarea=True)
    add_question(doc, "What industry or niche are you in? *", "Examples: ecommerce, coaching, consulting, real estate, SaaS, professional services.")
    add_question(doc, "How long have you been operating?")
    add_question(doc, "What stage is your business currently in? *", "Startup, early growth, scaling, established, restructuring, launching something new.")
    add_question(doc, "What do you currently sell or offer? *", "List your products, services, packages, memberships, retainers, programs, or digital offers.", is_textarea=True)
    add_question(doc, "What are your primary revenue streams?", "Examples: product sales, service retainers, project work, subscriptions, consulting, coaching, affiliate.")
    add_question(doc, "What is your current pricing structure?", "Include price ranges, retainers, one-time fees, subscriptions, or average order value if applicable.")
    add_question(doc, "What makes your business different from competitors?", is_textarea=True)

    # Section 3
    add_section_header(doc, "Section 3: Goals & Vision")
    add_question(doc, "What are your top 3 goals for the next 6–12 months? *", "Goals 1, 2, 3 details.", is_textarea=True)
    add_question(doc, "What is your primary goal from working with Sienvi Agency? *", "Example: increase leads, improve conversions, organize operations, launch a campaign, build systems, save time.")
    add_question(doc, "Specific SMART goal details: *", "What exactly do you want to achieve?", is_textarea=True)
    add_question(doc, "Measurable SMART goal metrics: *", "What numbers, KPIs, or results define success?", is_textarea=True)
    add_question(doc, "Achievable SMART goal parameters: *", "What resources, team, or budget do you already have?", is_textarea=True)
    add_question(doc, "Relevant SMART goal parameters: *", "Why does this goal matter right now?", is_textarea=True)
    add_question(doc, "Time-bound SMART goal deadline: *", "What is your deadline or ideal timeline?")
    add_question(doc, "What would a \"big win\" look like? *", "e.g., 100 leads/mo, 3x ROAS, 20% growth, 10 hours/week saved.", is_textarea=True)
    add_question(doc, "What is your bigger 3–5 year vision for the business?")
    add_question(doc, "Are there any major launches, campaigns, events, or deadlines coming up?")

    # Section 4
    add_section_header(doc, "Section 4: Current Challenges & Bottlenecks")
    add_question(doc, "What are the biggest challenges in your business right now? *", "Marketing, sales, operations, branding, content, admin, systems, etc.", is_textarea=True)
    add_question(doc, "Where do you feel stuck or overwhelmed? *", is_textarea=True)
    add_question(doc, "What is preventing you from reaching your goals faster?")
    add_question(doc, "What tasks or responsibilities are taking too much of your time?")
    add_question(doc, "What is currently inconsistent or not working well?")
    add_question(doc, "Worked with agencies/freelancers before? What worked and didn't?", is_textarea=True)
    add_question(doc, "What do you absolutely want to avoid in this project or partnership?", is_textarea=True)

    # Section 5
    add_section_header(doc, "Section 5: Target Audience & Customer Profile")
    add_question(doc, "Who is your ideal customer or client? *", "Describe the type of person or business you want to attract.", is_textarea=True)
    add_question(doc, "Is your audience B2C, B2B, or both?", None, options=["B2C", "B2B", "Both"])
    add_question(doc, "What demographics define your audience?", "Age, gender, location, income level, profession, business size.")
    add_question(doc, "What are their biggest goals or desires?")
    add_question(doc, "What problems, frustrations, or pain points do they have? *", is_textarea=True)
    add_question(doc, "What are they afraid will happen if they do not solve the problem?")
    add_question(doc, "What objections do they usually have?", "Price, trust, time, complexity, uncertainty.")
    add_question(doc, "Where does your audience spend time online?", "Google, Instagram, LinkedIn, TikTok, YouTube, Reddit, podcasts...")
    add_question(doc, "How do customers usually find you right now?")
    add_question(doc, "Which customer type is most valuable to your business? *", "Explain why: highest profit, easiest to close, best fit...", is_textarea=True)
    add_question(doc, "Customer types you want to avoid?", "Red flags, poor fit, wrong niche.")

    # Section 6
    add_section_header(doc, "Section 6: Offers, Products, Services & Positioning")
    add_question(doc, "What are your core offers, products, or services? *", "Short description of each.", is_textarea=True)
    add_question(doc, "Which offer is most important to grow right now? *")
    add_question(doc, "Offer segmentation (Most profitable vs. Easiest to sell vs. Needing improvement)?", is_textarea=True)
    add_question(doc, "What transformation does your customer get?", "Describe the \"before\" and \"after\" state.", is_textarea=True)
    add_question(doc, "What are your strongest selling points or unique advantages?")
    add_question(doc, "What proof do you have that your offer works?", "Testimonials, case studies, reviews, data...", is_textarea=True)
    add_question(doc, "Guarantees, warranties, bonuses, or incentives?")
    add_question(doc, "Are there claims, phrases, or promises we should avoid?")

    # Section 7
    add_section_header(doc, "Section 7: Marketing & Sales")
    add_question(doc, "How are you currently acquiring leads? *", "Paid ads, organic, SEO, referrals, outreach...", is_textarea=True)
    add_question(doc, "Marketing performance (What's working vs. inconsistent)?", is_textarea=True)
    add_question(doc, "Do you currently run paid advertising? *", None, options=["Yes", "No"])
    add_question(doc, "If yes, list platforms:", "Meta, Google, TikTok, LinkedIn, Amazon, YouTube...")
    add_question(doc, "Funnels, landing pages, or sequences active?", is_textarea=True)
    add_question(doc, "Current sales process & drop-off points?", "Inquiry -> Call -> Proposal -> Followup -> Close...", is_textarea=True)
    add_question(doc, "Who are your top 3 competitors?", "Include websites or social links.")
    add_question(doc, "How do you want to be different from competitors?")

    # Section 8
    add_section_header(doc, "Section 8: Brand, Messaging & Creative Direction")
    add_question(doc, "Do you have a clear brand identity? *", None, options=["Yes", "No"])
    add_question(doc, "Describe your brand voice: *", "e.g., professional, friendly, bold, minimal, playful, technical...")
    add_question(doc, "Associated words vs. Words to avoid?", "Words associated with brand vs. terms to avoid...", is_textarea=True)
    add_question(doc, "Admired brands & existing assets?", "Inspiring brands, photos, logos, templates...", is_textarea=True)
    add_question(doc, "Do you need help with design, copy, content, ads, or strategy?", is_textarea=True)

    # Section 9
    add_section_header(doc, "Section 9: Operations, Admin & Systems")
    add_question(doc, "List tools used (PM, CRM, Scheduling, Comm, File Storage, Invoicing, Zapier)", "ClickUp, HubSpot, Calendly, Slack, Drive, Stripe, Zapier...", is_textarea=True)
    add_question(doc, "What systems feel messy or hard to manage?", is_textarea=True)
    add_question(doc, "Documented SOPs or workflows?", None, options=["Yes", "No", "Partially"])
    add_question(doc, "Recurring tasks or manual automation needs?", is_textarea=True)

    # Section 10
    add_section_header(doc, "Section 10: Project Scope & Support Needed")
    add_question(doc, "Fulfillment support requested (Select all that apply): *", None, options=[
        "Business/Marketing strategy", "Paid advertising", "Social media & content",
        "Design & Website support", "Copywriting & Email marketing", "CRM & Automations",
        "Admin & EA support", "Ops cleanup", "Listing optimization"
    ])
    add_question(doc, "Top 3 priorities Sienvi should focus on first: *", "1. ...\n2. ...\n3. ...", is_textarea=True)
    add_question(doc, "First 7 days & first 30 days goals?", is_textarea=True)

    # Section 11
    add_section_header(doc, "Section 11: Assets, Access & Links")
    add_question(doc, "Please provide URLs (Website, Shared folder, brand assets, logos, CRM, PM):", is_textarea=True)
    add_question(doc, "Required platform access & access methods: *", "Direct invite, Password manager details...", is_textarea=True)
    add_question(doc, "Password Manager Used", "1Password, LastPass, Dashlane, none...")

    # Section 12
    add_section_header(doc, "Section 12: Communication & Collaboration")
    add_question(doc, "Collaboration preference & update frequency: *", "Slack async, weekly zoom, frequency...", is_textarea=False)
    add_question(doc, "Who approves work before it goes live? *")
    add_question(doc, "Feedback turnaround expectations & frustrations?", is_textarea=True)

    # Section 13
    add_section_header(doc, "Section 13: Timeline, Budget & Expectations")
    add_question(doc, "Start date, deadlines, project timeline & budget: *", "Expected start, budget ranges, deadlines...", is_textarea=True)

    # Section 14
    add_section_header(doc, "Section 14: Final Notes")
    add_question(doc, "Excitement vs. Worries & final context?", is_textarea=True)

    # Internal Checklist
    add_section_header(doc, "Internal Agency Review Checklist (Agency Use Only)")
    add_question(doc, "Client Type & Main Bottleneck", None, options=["Ecommerce", "Service", "Coaching", "SaaS", "Local", "Mixed"])
    add_question(doc, "Clarity Grade (Audience, Offer, Brand, Systems)", "Grade as Strong, Moderate, or Weak...", is_textarea=True)
    add_question(doc, "First 7/30 days recommendations & Upsells?", is_textarea=True)

    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")

if __name__ == "__main__":
    docs_dir = "c:\\Users\\Iris\\OneDrive\\Work\\sienvi-agency-landing-page\\docs"
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)
        
    build_business_admin_docx(os.path.join(docs_dir, "Business_Admin_Onboarding_Questionnaire_v2.docx"))
    build_general_discovery_docx(os.path.join(docs_dir, "General_Client_Onboarding_Discovery_Questionnaire_v2.docx"))
    print("DOCX generation complete!")
